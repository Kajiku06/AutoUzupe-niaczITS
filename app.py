from pyexpat.errors import codes
from flask import Flask, render_template, request, send_from_directory, url_for, jsonify
from docx import Document
import os
import re
import string
import threading
import time
from datetime import datetime
import json

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'output'

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

KEYS_FILE = os.path.join(os.path.dirname(__file__), 'keys_data.json')

def load_keys_from_file():
    if os.path.exists(KEYS_FILE):
        with open(KEYS_FILE, encoding='utf-8') as f:
            data = json.load(f)
            # Wczytaj wszystkie KEYS_* jako globalne zmienne
            for k, v in data.items():
                if k.startswith('KEYS_'):
                    globals()[k] = v
    else:
        save_keys_to_file()

def save_keys_to_file():
    data = {}
    # Zapisz tylko KEYS_* z globali
    for k, v in globals().items():
        if k.startswith('KEYS_'):
            data[k] = v
    with open(KEYS_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# Załaduj KEYS przy starcie (i utwórz plik jeśli nie istnieje)
load_keys_from_file()

def get_keys_by_name(name):
    # Zwraca tablicę KEYS_* z globali, zawsze listę
    return list(globals().get(name, []))

def parse_data(text, keys):
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    data = {}
    for i, key in enumerate(keys):
        if i >= len(lines):
            break
        line = lines[i]
        value = line.split(':', 1)[1].strip() if ':' in line else line
        data[key.lower()] = value
    return data

def replace_placeholders_in_paragraph(paragraph, data_dict):
    text = ''.join(run.text for run in paragraph.runs)
    original_text = text
    for key, value in data_dict.items():
        for placeholder in (f'[{key}]', f'{{{key}}}'):
            pattern = re.compile(re.escape(placeholder), re.IGNORECASE)
            text = pattern.sub(value, text)
    if text != original_text:
        for run in paragraph.runs:
            run.text = ''
        paragraph.runs[0].text = text

def replace_placeholders_in_cell(cell, data_dict):
    for paragraph in cell.paragraphs:
        replace_placeholders_in_paragraph(paragraph, data_dict)

def fill_word_template(template_path, data_dict, output_path):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, data_dict)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholders_in_cell(cell, data_dict)
    doc.save(output_path)

def sanitize_filename(filename):
    valid_chars = "-_.() %s%s" % (string.ascii_letters, string.digits)
    cleaned = ''.join(c for c in filename if c in valid_chars)
    return cleaned

def generate_output_filename(original_filename, parsed_data):
    boat_name_raw = parsed_data.get('boatname', 'output')
    boat_name = sanitize_filename(boat_name_raw.replace(' ', '_'))
    filename_lower = original_filename.lower()
    if 'mmsi' in filename_lower:
        suffix = "_MMSI"
    elif 'registration' in filename_lower:
        suffix = "_registration"
    else:
        suffix = ""
    final_name = f"{boat_name}{suffix}.docx"
    return final_name

def cleanup_old_files(folder, max_age_seconds=300):
    now = time.time()
    for filename in os.listdir(folder):
        filepath = os.path.join(folder, filename)
        if os.path.isfile(filepath):
            file_age = now - os.path.getmtime(filepath)
            if file_age > max_age_seconds:
                try:
                    os.remove(filepath)
                    print(f"Usunięto plik: {filepath}")
                except Exception as e:
                    print(f"Błąd podczas usuwania {filepath}: {e}")

def delayed_cleanup():
    time.sleep(300)
    cleanup_old_files(app.config['UPLOAD_FOLDER'])
    cleanup_old_files(app.config['OUTPUT_FOLDER'])

def is_valid_hin(val):
    if len(val) != 15:
        return False
    allowed = string.ascii_letters + string.digits
    return all(c in allowed for c in val[0:2]) and val[2] == '-' and all(c in allowed for c in val[3:])

def extract_alphanum_start(s):
    for i, c in enumerate(s):
        if c.isalnum():
            return s[i:]
    return s

def filter_hin_and_category_keys(keys, raw_text):
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    filtered_keys = []
    for idx, key in enumerate(keys):
        key_lower = key.lower()
        if key_lower == 'hinnumber':
            if idx >= len(lines):
                continue
            val = lines[idx].split(':', 1)[-1]
            val = extract_alphanum_start(val.strip())
            if is_valid_hin(val):
                filtered_keys.append(key)
            else:
                filtered_keys.append(key)
        elif key_lower == 'category':
            if idx >= len(lines):
                continue
            val = lines[idx].split(':', 1)[-1]
            val = extract_alphanum_start(val.strip()).upper()
            if val in ['A', 'B', 'C', 'D']:
                filtered_keys.append(key)
        else:
            filtered_keys.append(key)
    return filtered_keys

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/parse', methods=['POST'])
def parse():
    raw_text = request.form.get('data', '')
    file1 = request.files.get('file1')
    file2 = request.files.get('file2')
    operation_type = request.form.get('operationType', 'rejestracja')
    people_count = request.form.get('peopleCount', '1')
    is_company = request.form.get('isCompany') == 'on'
    second_engine = request.form.get('secondEngine') == 'on'
    no_file = request.form.get('noFile') == 'on'

    def pick_keys():
        # Nazwy tablic KEYS_* zgodnie z logiką
        if operation_type == 'rejestracja':
            if people_count == '1':
                if is_company and second_engine:
                    return get_keys_by_name('KEYS_REJ_1_FIRMA_2ENG')
                elif is_company:
                    return get_keys_by_name('KEYS_REJ_1_FIRMA')
                elif second_engine:
                    return get_keys_by_name('KEYS_REJ_1_2ENG')
                else:
                    return get_keys_by_name('KEYS_REJ_1')
            elif people_count == '2':
                if is_company and second_engine:
                    return get_keys_by_name('KEYS_REJ_2_FIRMA_2ENG')
                elif is_company:
                    return get_keys_by_name('KEYS_REJ_2_FIRMA')
                elif second_engine:
                    return get_keys_by_name('KEYS_REJ_2_2ENG')
                else:
                    return get_keys_by_name('KEYS_REJ_2')
            elif people_count == '3':
                if is_company and second_engine:
                    return get_keys_by_name('KEYS_REJ_3_FIRMA_2ENG')
                elif is_company:
                    return get_keys_by_name('KEYS_REJ_3_FIRMA')
                elif second_engine:
                    return get_keys_by_name('KEYS_REJ_3_2ENG')
                else:
                    return get_keys_by_name('KEYS_REJ_3')
            else:
                return get_keys_by_name('KEYS_REJ_1')
        elif operation_type == 'zmiana':
            if people_count == '1':
                if is_company and second_engine:
                    return get_keys_by_name('KEYS_ZMIANA_1_FIRMA')
                elif is_company:
                    return get_keys_by_name('KEYS_ZMIANA_1_FIRMA')
                elif second_engine:
                    return get_keys_by_name('KEYS_ZMIANA_1')
                else:
                    return get_keys_by_name('KEYS_ZMIANA_1')
            elif people_count == '2':
                if is_company and second_engine:
                    return get_keys_by_name('KEYS_ZMIANA_2_FIRMA')
                elif is_company:
                    return get_keys_by_name('KEYS_ZMIANA_2_FIRMA')
                elif second_engine:
                    return get_keys_by_name('KEYS_ZMIANA_2')
                else:
                    return get_keys_by_name('KEYS_ZMIANA_2')
            elif people_count == '3':
                if is_company and second_engine:
                    return get_keys_by_name('KEYS_ZMIANA_3_FIRMA')
                elif is_company:
                    return get_keys_by_name('KEYS_ZMIANA_3_FIRMA')
                elif second_engine:
                    return get_keys_by_name('KEYS_ZMIANA_3')
                else:
                    return get_keys_by_name('KEYS_ZMIANA_3')
            else:
                return get_keys_by_name('KEYS_ZMIANA_1')
        else:
            return get_keys_by_name('KEYS_REJ_1')

    keys = pick_keys()
    keys_filtered = filter_hin_and_category_keys(keys, raw_text)
    parsed_data = parse_data(raw_text, keys_filtered)
    parsed_data['todaydata'] = datetime.now().strftime("%Y-%m-%d")

    # Nadpisz HINNumber na '-----' jeśli niepoprawny (ale nie dodawaj na końcu)
    for idx, key in enumerate(keys_filtered):
        if key.lower() == 'hinnumber':
            val = parsed_data.get('hinnumber', '')
            if not is_valid_hin(val):
                parsed_data['hinnumber'] = '-----'

    # Dodaj category na końcu jeśli nie ma w keys_filtered
    keys_lower = [k.lower() for k in keys_filtered]
    if 'category' not in keys_lower:
        parsed_data['category'] = '-----'

    files_to_process = []
    if not no_file:
        if file1 and file1.filename.lower().endswith('.docx'):
            files_to_process.append(('plik1', file1))
        if file2 and file2.filename.lower().endswith('.docx'):
            files_to_process.append(('plik2', file2))

        if not files_to_process:
            return render_template('result.html', data=parsed_data, output_files=[], no_file=True, error="Nie wybrano żadnego pliku do uzupełnienia. Zaznacz 'Tylko parsuj dane' jeśli chcesz tylko sprawdzić dane.")

    output_files = []
    for label, file in files_to_process:
        template_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(template_path)
        output_filename = generate_output_filename(file.filename, parsed_data)
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        fill_word_template(template_path, parsed_data, output_path)
        output_files.append((label, output_filename))

    threading.Thread(target=delayed_cleanup, daemon=True).start()

    return render_template(
        'result.html',
        data=parsed_data,
        output_files=output_files,
        no_file=no_file,
        error=None
    )

@app.route('/download/<filename>')
def download(filename):
    return send_from_directory(app.config['OUTPUT_FOLDER'], filename, as_attachment=True)

@app.route('/instructions')
def instructions():
    return render_template('instructions.html')

@app.route('/football')
def football():
    return render_template('football.html')

@app.route('/football_score', methods=['POST'])
def football_score():
    data = request.json
    name = data.get('name', 'anonim')
    score = data.get('score', 0)
    # Zapisz wynik do pliku tekstowego (prosty sposób)
    with open('football_scores.txt', 'a', encoding='utf-8') as f:
        f.write(f"{name}:{score}\n")
    return jsonify({'status': 'ok'})

@app.route('/football_scoreboard')
def football_scoreboard():
    scores = []
    try:
        with open('football_scores.txt', encoding='utf-8') as f:
            for line in f:
                if ':' in line:
                    # Rozdziel po ostatnim dwukropku
                    parts = line.strip().rsplit(':', 1)
                    if len(parts) == 2:
                        name, score = parts
                        try:
                            scores.append({'name': name, 'score': int(score)})
                        except ValueError:
                            continue
        # Sortuj malejąco po wyniku, top 3
        scores = sorted(scores, key=lambda x: x['score'], reverse=True)[:3]
    except Exception:
        scores = []
    return jsonify(scores)

@app.route('/last_update')
def last_update():
    try:
        filepath = os.path.join(os.path.dirname(__file__), 'app.py')
        mtime = os.path.getmtime(filepath)
        dt = datetime.fromtimestamp(mtime)
        # Format: YYYY-MM-DD HH:MM:SS
        return jsonify({'last_update': dt.strftime('%Y-%m-%d %H:%M:%S')})
    except Exception:
        return jsonify({'last_update': 'brak danych'})

# --- Panel admina: edycja KEYS ---

KEYS_FILE = os.path.join(os.path.dirname(__file__), 'keys_data.json')

def load_keys_from_file():
    if os.path.exists(KEYS_FILE):
        with open(KEYS_FILE, encoding='utf-8') as f:
            data = json.load(f)
            for k, v in data.items():
                if k.startswith('KEYS_'):
                    globals()[k] = v
    else:
        # Jeśli plik nie istnieje, utwórz go z aktualnymi KEYS_* z kodu
        save_keys_to_file()

def save_keys_to_file():
    data = {}
    for k, v in globals().items():
        if k.startswith('KEYS_'):
            data[k] = v
    with open(KEYS_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# Załaduj KEYS przy starcie (i utwórz plik jeśli nie istnieje)
load_keys_from_file()

@app.route('/admin')
def admin_login():
    return render_template('admin_login.html')

@app.route('/admin/keys')
def admin_keys():
    return render_template('admin_keys.html')

@app.route('/get_keys')
def get_keys():
    all_keys = request.args.get('all')
    result = {}
    if all_keys:
        # Zwróć wszystkie KEYS_* z globali jako listy (nie numpy arrays, nie None)
        for k, v in globals().items():
            if k.startswith('KEYS_'):
                # Upewnij się, że to lista (może być tuple po *KEYS_ZMIANA_3_FIRMA, ...)
                if isinstance(v, (list, tuple)):
                    result[k] = list(v)
                else:
                    result[k] = []
        return jsonify(result)
    # Domyślnie tylko wybrane (do zachowania kompatybilności)
    for name in [
        "KEYS_REJ_2_FIRMA",
        "KEYS_REJ_2_FIRMA_2ENG",
        "KEYS_REJ_3",
        "KEYS_REJ_3_2ENG",
        "KEYS_REJ_3_FIRMA",
        "KEYS_REJ_3_FIRMA_2ENG"
    ]:
        v = globals().get(name, [])
        result[name] = list(v) if isinstance(v, (list, tuple)) else []
    return jsonify(result)

@app.route('/set_key', methods=['POST'])
def set_key():
    data = request.get_json()
    key = data.get('key')
    value = data.get('value', '')
    if key and key.startswith('KEYS_'):
        arr = [x.strip() for x in value.split(',') if x.strip()]
        globals()[key] = arr
        save_keys_to_file()
        return jsonify({'status': 'ok'})
    return jsonify({'status': 'error'})

# Wszelkie zmiany przez panel admina są zapisywane do keys_data.json przez save_keys_to_file()
# KEYS_* są trwale przechowywane w tym pliku i nie są kasowane automatycznie.
# Usunięcie lub zmiana KEYS_* następuje tylko przez panel admina lub ręczną edycję pliku keys_data.json.

if __name__ == '__main__':
    app.run(debug=True)

